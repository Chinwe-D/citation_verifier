#!/usr/bin/env python3
"""
link_citations.py

Step 5 of the citation-verification pipeline: produces a copy of the
manuscript where every in-text citation is an internal hyperlink to its
entry in the reference list (click "Adiele et al. (2020)" in the body,
land on the Adiele, J.G. et al. (2020) paragraph in References).

Approach:
  1. Bookmark every reference-list paragraph (one bookmark per entry,
     named from first-author-surname + year).
  2. Scan body paragraphs for the same citation patterns used elsewhere
     in this pipeline. For each match found *entirely within a single
     run* (the common case after Word's own run-merging, or after
     running the docx skill's merge_runs.py on messier files), split
     that run and wrap the citation text in an internal hyperlink
     pointing at the matching bookmark.
  3. Citations that fall across multiple runs (rare, but possible with
     heavily-edited/tracked-changes documents) are left unlinked and
     reported at the end -- link those by hand, or run merge_runs.py on
     the source file first (see docx skill) to reduce run fragmentation.

Usage:
    python link_citations.py manuscript.docx output.docx
"""
import sys
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, __file__.rsplit("/", 1)[0])
from extract_citations import detect_citation_style


def make_bookmark_start(bookmark_id, name):
    el = OxmlElement("w:bookmarkStart")
    el.set(qn("w:id"), str(bookmark_id))
    el.set(qn("w:name"), name)
    return el


def make_bookmark_end(bookmark_id):
    el = OxmlElement("w:bookmarkEnd")
    el.set(qn("w:id"), str(bookmark_id))
    return el


def make_hyperlink_run(anchor_name, run_text, rpr_source_run):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor_name)

    new_run = OxmlElement("w:r")
    if rpr_source_run is not None and rpr_source_run.find(qn("w:rPr")) is not None:
        rpr = copy.deepcopy(rpr_source_run.find(qn("w:rPr")))
        new_run.append(rpr)
    # Style citation links: underline + a distinct color, in addition to
    # whatever formatting the original run had.
    rpr = new_run.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        new_run.insert(0, rpr)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = run_text
    new_run.append(t)
    hyperlink.append(new_run)
    return hyperlink


def slugify(surname, year):
    base_year = re.match(r"(\d{4})", year).group(1)
    return f"ref_{re.sub(r'[^A-Za-z0-9]', '', surname)}_{base_year}"


REF_HEADING_PAT = re.compile(r"^\s*References?\s*$", re.IGNORECASE)
REF_ENTRY_PAT = re.compile(r"^([A-Z][A-Za-zÀ-ÿ\-']+)[,.\s].*?\((\d{4}[a-z]?(?:/\d{4})?)\)")
NUMBERED_REF_ENTRY_PAT = re.compile(r"^\s*\[?(\d{1,3})\]?[\.\):]?\s+\S")
NUMBERED_INTEXT_PAT = re.compile(r"[\[\(]\s*(\d{1,3}(?:\s*[-,–]\s*\d{1,3})*)\s*[\]\)]")

NARRATIVE_PAT = re.compile(
    r"\b([A-Z][A-Za-zÀ-ÿ\-']+)(?:\s+(?:et al\.|and\s+[A-Z][A-Za-zÀ-ÿ\-']+))?\s*"
    r"\((\d{4}[a-z]?(?:/\d{4})?)\)"
)
PARENTHETICAL_ENTRY_PAT = re.compile(
    r"([A-Z][A-Za-zÀ-ÿ\-']+)(?:\s+et al\.)?,\s*(\d{4}[a-z]?)"
)


def find_references_section(doc):
    for i, p in enumerate(doc.paragraphs):
        if REF_HEADING_PAT.match(p.text.strip()):
            return i
    return None


def bookmark_references(doc, ref_start_idx, ambiguity_report):
    """
    Add a bookmark to each reference-list paragraph.
    Returns dict {(surname, base_year): [(anchor, n_authors), ...]} -- a list
    because two references can legitimately share the same first-author +
    year (e.g. a PhD thesis and a related journal paper), which also means
    the manuscript's citation style should disambiguate them as '2020a'/'2020b'.
    """
    bookmark_id = 1000
    ref_keys = {}
    for p in doc.paragraphs[ref_start_idx + 1:]:
        text = p.text.strip()
        if not text:
            continue
        m = REF_ENTRY_PAT.match(text)
        if not m:
            continue
        surname, year = m.group(1), m.group(2)
        base_year = re.match(r"(\d{4})", year).group(1)
        n_authors = len(re.findall(r"[A-Z][A-Za-zÀ-ÿ\-']+,\s*[A-Z]\.", text.split("(")[0])) or 1

        key = (surname, base_year)
        n_existing = len(ref_keys.get(key, []))
        anchor = slugify(surname, year) + (f"_{n_existing+1}" if n_existing else "")
        ref_keys.setdefault(key, []).append((anchor, n_authors))

        if n_existing == 1:  # this is the 2nd entry sharing the key
            ambiguity_report.append(
                f"'{surname} ({year})' is used by more than one reference-list entry "
                f"-- consider disambiguating as {year}a / {year}b in both the text and "
                f"the reference list, per most citation styles."
            )

        first_run_el = p.runs[0]._r if p.runs else None
        start_el = make_bookmark_start(bookmark_id, anchor)
        end_el = make_bookmark_end(bookmark_id)
        if first_run_el is not None:
            first_run_el.addprevious(start_el)
            p._p.append(end_el)
        else:
            p._p.insert(0, start_el)
            p._p.append(end_el)
        bookmark_id += 1
    return ref_keys


def resolve_anchor(ref_keys, surname, base_year, cited_as_et_al):
    """Pick the right anchor when a (surname, year) key has multiple entries."""
    candidates = ref_keys.get((surname, base_year))
    if not candidates:
        return None
    if len(candidates) == 1:
        return candidates[0][0]
    # Prefer a multi-author entry for 'et al.' citations, a single-author
    # entry for bare 'Surname (year)' citations; fall back to the first.
    for anchor, n_authors in candidates:
        if cited_as_et_al and n_authors > 1:
            return anchor
        if not cited_as_et_al and n_authors == 1:
            return anchor
    return candidates[0][0]


def link_citations_in_paragraph(p, ref_keys, unlinked_report):
    """Find citations in this paragraph and hyperlink any that fall within a single run."""
    # Build run text + cumulative offsets
    runs = p.runs
    if not runs:
        return
    offsets = []
    pos = 0
    for r in runs:
        offsets.append((pos, pos + len(r.text), r))
        pos += len(r.text)
    full_text = "".join(r.text for r in runs)

    matches = []
    for m in NARRATIVE_PAT.finditer(full_text):
        surname, year = m.group(1), m.group(2)
        base_year = re.match(r"(\d{4})", year).group(1)
        if (surname, base_year) in ref_keys:
            cited_as_et_al = "et al" in m.group(0)
            matches.append((m.start(), m.end(), surname, year, cited_as_et_al))
    # Parenthetical citations -- link just the "Surname, YYYY" substring inside the parens
    for pm in re.finditer(r"\(([^()]*\d{4}[^()]*)\)", full_text):
        inner = pm.group(1)
        for em in PARENTHETICAL_ENTRY_PAT.finditer(inner):
            surname, year = em.group(1), em.group(2)
            base_year = re.match(r"(\d{4})", year).group(1)
            if (surname, base_year) in ref_keys:
                abs_start = pm.start(1) + em.start()
                abs_end = pm.start(1) + em.end()
                cited_as_et_al = "et al" in em.group(0)
                matches.append((abs_start, abs_end, surname, year, cited_as_et_al))

    if not matches:
        return

    # Group matches by the run they fall entirely within, so a run with
    # multiple citations is rebuilt exactly once (rebuilding per-match would
    # detach an already-replaced run element on the second match).
    by_run = {}
    for start, end, surname, year, cited_as_et_al in matches:
        target = None
        for r_start, r_end, run in offsets:
            if r_start <= start and end <= r_end:
                target = (r_start, r_end, run)
                break
        if target is None:
            unlinked_report.append(f"{surname} ({year}) -- spans multiple runs, left unlinked")
            continue
        by_run.setdefault(id(target[2]), (target, []))[1].append((start, end, surname, year, cited_as_et_al))

    for (r_start, r_end, run), run_matches in by_run.values():
        run_matches.sort(key=lambda x: x[0])  # left to right within this run
        full_run_text = run.text
        run_el = run._r
        parent = run_el.getparent()
        idx = list(parent).index(run_el)

        new_elements = []
        cursor = 0  # local offset within this run's text
        for start, end, surname, year, cited_as_et_al in run_matches:
            local_start, local_end = start - r_start, end - r_start
            before_text = full_run_text[cursor:local_start]
            citation_text = full_run_text[local_start:local_end]
            if before_text:
                before_run = copy.deepcopy(run_el)
                t = before_run.find(qn("w:t"))
                t.text = before_text
                t.set(qn("xml:space"), "preserve")
                new_elements.append(before_run)
            base_year = re.match(r"(\d{4})", year).group(1)
            anchor = resolve_anchor(ref_keys, surname, base_year, cited_as_et_al)
            new_elements.append(make_hyperlink_run(anchor, citation_text, run_el))
            cursor = local_end

        tail_text = full_run_text[cursor:]
        if tail_text:
            after_run = copy.deepcopy(run_el)
            t = after_run.find(qn("w:t"))
            t.text = tail_text
            t.set(qn("xml:space"), "preserve")
            new_elements.append(after_run)

        parent.remove(run_el)
        for i, el in enumerate(new_elements):
            parent.insert(idx + i, el)


def bookmark_references_numbered(doc, ref_start_idx):
    """Bookmark each numbered reference-list paragraph. Returns {number: anchor}."""
    bookmark_id = 2000
    ref_numbers = {}
    for p in doc.paragraphs[ref_start_idx + 1:]:
        text = p.text.strip()
        if not text:
            continue
        m = NUMBERED_REF_ENTRY_PAT.match(text)
        if not m:
            continue
        number = int(m.group(1))
        anchor = f"ref_num_{number}"
        ref_numbers[number] = anchor

        first_run_el = p.runs[0]._r if p.runs else None
        start_el = make_bookmark_start(bookmark_id, anchor)
        end_el = make_bookmark_end(bookmark_id)
        if first_run_el is not None:
            first_run_el.addprevious(start_el)
            p._p.append(end_el)
        else:
            p._p.insert(0, start_el)
            p._p.append(end_el)
        bookmark_id += 1
    return ref_numbers


def link_citations_in_paragraph_numbered(p, ref_numbers, unlinked_report):
    """
    Hyperlink [1], [2,3], [4-6] style citations. Since a single bracket
    group can reference several numbers at once (e.g. '[4-6]'), and Word
    hyperlinks can only point at one bookmark, a multi-number group links
    to its *first* valid number -- the group still reads correctly and
    lands the reader in the reference list, just not on every number in a
    range individually.
    """
    runs = p.runs
    if not runs:
        return
    offsets = []
    pos = 0
    for r in runs:
        offsets.append((pos, pos + len(r.text), r))
        pos += len(r.text)
    full_text = "".join(r.text for r in runs)

    matches = []
    for m in NUMBERED_INTEXT_PAT.finditer(full_text):
        group = m.group(1)
        numbers = []
        for part in re.split(r",\s*", group):
            part = part.strip()
            if "-" in part or "–" in part:
                a, b = re.split(r"[-–]", part)
                a, b = int(a), int(b)
                if 0 < b - a < 50:
                    numbers.extend(range(a, b + 1))
            elif part.isdigit():
                numbers.append(int(part))
        valid = [n for n in numbers if n in ref_numbers]
        if not valid:
            continue
        matches.append((m.start(), m.end(), valid[0], group))

    if not matches:
        return

    by_run = {}
    for start, end, number, group in matches:
        target = None
        for r_start, r_end, run in offsets:
            if r_start <= start and end <= r_end:
                target = (r_start, r_end, run)
                break
        if target is None:
            unlinked_report.append(f"[{group}] -- spans multiple runs, left unlinked")
            continue
        by_run.setdefault(id(target[2]), (target, []))[1].append((start, end, number, group))

    for (r_start, r_end, run), run_matches in by_run.values():
        run_matches.sort(key=lambda x: x[0])
        full_run_text = run.text
        run_el = run._r
        parent = run_el.getparent()
        idx = list(parent).index(run_el)

        new_elements = []
        cursor = 0
        for start, end, number, group in run_matches:
            local_start, local_end = start - r_start, end - r_start
            before_text = full_run_text[cursor:local_start]
            citation_text = full_run_text[local_start:local_end]
            if before_text:
                before_run = copy.deepcopy(run_el)
                t = before_run.find(qn("w:t"))
                t.text = before_text
                t.set(qn("xml:space"), "preserve")
                new_elements.append(before_run)
            new_elements.append(make_hyperlink_run(ref_numbers[number], citation_text, run_el))
            cursor = local_end

        tail_text = full_run_text[cursor:]
        if tail_text:
            after_run = copy.deepcopy(run_el)
            t = after_run.find(qn("w:t"))
            t.text = tail_text
            t.set(qn("xml:space"), "preserve")
            new_elements.append(after_run)

        parent.remove(run_el)
        for i, el in enumerate(new_elements):
            parent.insert(idx + i, el)


def link_manuscript(doc, ref_start_idx, style="auto"):
    """
    Style-aware linking entry point used by the app. Returns
    (resolved_style, ambiguity_report, unlinked_report).
    """
    ref_text_block = "\n\n".join(p.text for p in doc.paragraphs[ref_start_idx + 1:])
    resolved_style = detect_citation_style(ref_text_block) if style == "auto" else style

    ambiguity_report, unlinked_report = [], []
    if resolved_style == "numbered":
        ref_numbers = bookmark_references_numbered(doc, ref_start_idx)
        for p in doc.paragraphs[:ref_start_idx]:
            link_citations_in_paragraph_numbered(p, ref_numbers, unlinked_report)
    else:
        resolved_style = "author_year"
        ref_keys = bookmark_references(doc, ref_start_idx, ambiguity_report)
        for p in doc.paragraphs[:ref_start_idx]:
            link_citations_in_paragraph(p, ref_keys, unlinked_report)

    return resolved_style, ambiguity_report, unlinked_report


def main():
    if len(sys.argv) not in (3, 4):
        print("Usage: python link_citations.py manuscript.docx output.docx [author_year|numbered|auto]")
        sys.exit(1)

    in_path, out_path = sys.argv[1], sys.argv[2]
    style = sys.argv[3] if len(sys.argv) == 4 else "auto"
    doc = Document(in_path)

    ref_start_idx = find_references_section(doc)
    if ref_start_idx is None:
        print("Could not find a 'References' heading paragraph -- aborting.")
        sys.exit(1)

    resolved_style, ambiguity_report, unlinked_report = link_manuscript(doc, ref_start_idx, style)
    print(f"Detected/using style: {resolved_style}")

    doc.save(out_path)
    print(f"Saved hyperlinked manuscript to {out_path}")

    if ambiguity_report:
        print(f"\n{len(ambiguity_report)} author-year collision(s) found:")
        for line in ambiguity_report:
            print(" -", line)
    if unlinked_report:
        print(f"\n{len(unlinked_report)} citation(s) could not be auto-linked (span multiple runs):")
        for line in unlinked_report:
            print(" -", line)


if __name__ == "__main__":
    main()
